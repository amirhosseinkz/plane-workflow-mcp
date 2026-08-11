from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from pypdf import PdfReader

import reports


class _Api:
    def project(self, project_id: str) -> dict[str, object]:
        return {"id": project_id, "identifier": "DEMO", "name": "Demo Project"}

    @staticmethod
    def states(project_id: str) -> list[dict[str, object]]:
        return [
            {"id": "backlog", "name": "Backlog", "type": "backlog"},
            {"id": "progress", "name": "In Progress", "type": "started"},
            {"id": "done", "name": "Done", "type": "completed"},
        ]

    @staticmethod
    def modules(project_id: str) -> list[dict[str, object]]:
        return [{"id": "mobile", "name": "Mobile"}]

    @staticmethod
    def labels(project_id: str) -> list[dict[str, object]]:
        return [{"id": "bug", "name": "Bug"}]

    @staticmethod
    def work_items(project_id: str) -> tuple[list[dict[str, object]], int]:
        return [
            {"id": "one", "sequence_id": 1, "name": "Fix login", "state": {"id": "backlog"}, "priority": "high", "labels": [{"id": "bug"}], "assignees": [{"member": {"id": "a", "display_name": "Ava"}}], "updated_at": "2026-08-01T08:30:00Z", "description_stripped": "Retry login safely."},
            {"id": "two", "sequence_id": 2, "name": "Improve search", "state": {"id": "progress"}, "priority": "medium", "labels": [], "assignees": [], "updated_at": "2026-08-02T08:30:00Z"},
            {"id": "three", "sequence_id": 3, "name": "Ship onboarding", "state": {"id": "done"}, "priority": "low", "labels": [], "assignees": [], "updated_at": "2026-08-03T08:30:00Z"},
        ], 3

    @staticmethod
    def module_work_items(project_id: str, module_id: str) -> list[dict[str, object]]:
        return [{"id": "one"}, {"id": "two"}]


class ReportTests(unittest.TestCase):
    def test_exports_filtered_docx_with_real_tables(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            result = reports.export_work_items_report(
                _Api(),
                project_id="project",
                report_format="docx",
                title="Backlog and in-progress work",
                filters={"state_names": ["Backlog", "In Progress"]},
                layout={"group_by": "state", "columns": ["identifier", "title", "state", "priority"]},
                output_directory=directory,
            )
            path = Path(result["path"])
            document = Document(path)

            self.assertTrue(path.exists())
            self.assertEqual(result["work_item_count"], 2)
            self.assertIn("Backlog and in-progress work", "\n".join(paragraph.text for paragraph in document.paragraphs))
            self.assertGreaterEqual(len(document.tables), 3)

    def test_exports_filtered_pdf_with_selected_items(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            result = reports.export_work_items_report(
                _Api(),
                project_id="project",
                report_format="pdf",
                title=None,
                filters={"include_completed": False},
                layout={"group_by": "state", "columns": ["identifier", "title", "state"]},
                output_directory=directory,
            )
            path = Path(result["path"])
            rendered = "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)

            self.assertTrue(path.exists())
            self.assertIn("Fix login", rendered)
            self.assertIn("Improve search", rendered)
            self.assertNotIn("Ship onboarding", rendered)
