"""Read-only Plane work-item report selection and rendering."""

from __future__ import annotations

import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Literal

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ReportFormat = Literal["docx", "pdf"]
VALID_GROUPS = {"state", "module", "priority", "none"}
VALID_COLUMNS = {"identifier", "title", "state", "priority", "assignees", "updated_at", "module", "labels", "description"}
DEFAULT_COLUMNS = ["identifier", "title", "state", "priority", "assignees", "updated_at"]


class ReportError(ValueError):
    """A safe, user-correctable report request error."""


@dataclass(frozen=True)
class ReportRow:
    identifier: str
    title: str
    state: str
    state_type: str
    priority: str
    assignees: str
    updated_at: str
    module: str
    labels: str
    description: str

    def value(self, column: str) -> str:
        return str(getattr(self, column))


@dataclass(frozen=True)
class ReportDocument:
    title: str
    project: dict[str, Any]
    generated_at: str
    filters: dict[str, Any]
    columns: list[str]
    group_by: str
    rows: list[ReportRow]
    total_work_items: int | None

    @property
    def groups(self) -> list[tuple[str, list[ReportRow]]]:
        grouped: dict[str, list[ReportRow]] = defaultdict(list)
        for row in self.rows:
            key = row.value(self.group_by) if self.group_by != "none" else "Work items"
            grouped[key or "Unspecified"].append(row)
        return [(name, grouped[name]) for name in sorted(grouped, key=lambda value: value.casefold())]


def _text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def _identifier(project: dict[str, Any], item: dict[str, Any]) -> str:
    prefix = _text(project.get("identifier"))
    sequence = item.get("sequence_id")
    return f"{prefix}-{sequence}" if prefix and sequence not in {None, ""} else _text(item.get("id"))


def _state_id(item: dict[str, Any]) -> str:
    state = item.get("state")
    if isinstance(state, dict):
        return _text(state.get("id"))
    return _text(state)


def _nested_ids(value: object, *, nested: str | None = None) -> set[str]:
    identifiers: set[str] = set()
    values = value if isinstance(value, list) else []
    for item in values:
        if isinstance(item, str):
            identifiers.add(item)
        elif isinstance(item, dict):
            candidate = item.get(nested) if nested else item
            if isinstance(candidate, dict) and candidate.get("id"):
                identifiers.add(str(candidate["id"]))
            elif item.get("id"):
                identifiers.add(str(item["id"]))
    return identifiers


def _display_names(value: object, *, nested: str | None = None) -> str:
    values = value if isinstance(value, list) else []
    names: list[str] = []
    for item in values:
        if isinstance(item, str):
            names.append(item)
        elif isinstance(item, dict):
            candidate = item.get(nested) if nested else item
            if isinstance(candidate, dict):
                names.append(_text(candidate.get("display_name") or candidate.get("name") or candidate.get("email")))
            else:
                names.append(_text(item.get("display_name") or item.get("name")))
    return ", ".join(name for name in names if name)


def _description(item: dict[str, Any]) -> str:
    raw = item.get("description_stripped") or item.get("description") or item.get("description_html") or ""
    return _text(re.sub(r"<[^>]+>", " ", str(raw)))


def _as_string_list(value: object, field: str) -> list[str]:
    if value is None:
        return []
    if not isinstance(value, list) or not all(isinstance(item, str) and _text(item) for item in value):
        raise ReportError(f"filters.{field} must be a list of nonempty strings.")
    return [_text(item) for item in value]


def _as_date(value: object, field: str) -> datetime | None:
    if value is None:
        return None
    if not isinstance(value, str):
        raise ReportError(f"filters.{field} must use YYYY-MM-DD.")
    try:
        return datetime.fromisoformat(value.strip()).replace(tzinfo=timezone.utc)
    except ValueError as error:
        raise ReportError(f"filters.{field} must use YYYY-MM-DD.") from error


def _validate_spec(filters: dict[str, Any] | None, layout: dict[str, Any] | None) -> tuple[dict[str, Any], dict[str, Any]]:
    if filters is not None and not isinstance(filters, dict):
        raise ReportError("filters must be an object.")
    if layout is not None and not isinstance(layout, dict):
        raise ReportError("layout must be an object.")
    normalized_filters = dict(filters or {})
    normalized_layout = dict(layout or {})
    supported_filters = {
        "state_names", "state_ids", "module_names", "module_ids", "label_names", "label_ids",
        "assignee_ids", "priorities", "include_completed", "updated_after", "updated_before",
    }
    unknown_filters = sorted(set(normalized_filters) - supported_filters)
    if unknown_filters:
        raise ReportError(f"Unsupported report filters: {', '.join(unknown_filters)}.")
    for field in supported_filters - {"include_completed", "updated_after", "updated_before"}:
        normalized_filters[field] = _as_string_list(normalized_filters.get(field), field)
    include_completed = normalized_filters.get("include_completed", True)
    if not isinstance(include_completed, bool):
        raise ReportError("filters.include_completed must be true or false.")
    normalized_filters["include_completed"] = include_completed
    normalized_filters["updated_after"] = _as_date(normalized_filters.get("updated_after"), "updated_after")
    normalized_filters["updated_before"] = _as_date(normalized_filters.get("updated_before"), "updated_before")
    if normalized_filters["updated_after"] and normalized_filters["updated_before"] and normalized_filters["updated_after"] > normalized_filters["updated_before"]:
        raise ReportError("filters.updated_after must not be later than filters.updated_before.")
    group_by = normalized_layout.get("group_by", "state")
    if group_by not in VALID_GROUPS:
        raise ReportError("layout.group_by must be state, module, priority, or none.")
    columns = normalized_layout.get("columns", DEFAULT_COLUMNS)
    if not isinstance(columns, list) or not columns or not all(isinstance(column, str) and column in VALID_COLUMNS for column in columns):
        raise ReportError(f"layout.columns must use supported columns: {', '.join(sorted(VALID_COLUMNS))}.")
    if len(columns) > 7:
        raise ReportError("layout.columns supports at most seven columns for readable reports.")
    normalized_layout["group_by"] = group_by
    normalized_layout["columns"] = columns
    return normalized_filters, normalized_layout


def build_report(
    api: Any,
    *,
    project_id: str,
    title: str | None,
    filters: dict[str, Any] | None,
    layout: dict[str, Any] | None,
) -> ReportDocument:
    normalized_filters, normalized_layout = _validate_spec(filters, layout)
    project = api.project(project_id)
    states = api.states(project_id)
    modules = api.modules(project_id)
    labels = api.labels(project_id)
    items, total_count = api.work_items(project_id)
    states_by_id = {str(item.get("id")): item for item in states if item.get("id")}
    modules_by_id = {str(item.get("id")): item for item in modules if item.get("id")}
    labels_by_id = {str(item.get("id")): item for item in labels if item.get("id")}

    def names_to_ids(names: list[str], collection: dict[str, dict[str, Any]], label: str) -> set[str]:
        wanted = {name.casefold() for name in names}
        found = {identifier for identifier, item in collection.items() if _text(item.get("name")).casefold() in wanted}
        missing = sorted(wanted - {_text(item.get("name")).casefold() for item in collection.values()})
        if missing:
            raise ReportError(f"Unknown {label}: {', '.join(missing)}.")
        return found

    state_ids = set(normalized_filters["state_ids"]) | names_to_ids(normalized_filters["state_names"], states_by_id, "state names")
    module_ids = set(normalized_filters["module_ids"]) | names_to_ids(normalized_filters["module_names"], modules_by_id, "module names")
    label_ids = set(normalized_filters["label_ids"]) | names_to_ids(normalized_filters["label_names"], labels_by_id, "label names")
    module_by_item_id: dict[str, set[str]] = defaultdict(set)
    if module_ids or normalized_layout["group_by"] == "module":
        for module_id in modules_by_id:
            for item in api.module_work_items(project_id, module_id):
                if item.get("id"):
                    module_by_item_id[str(item["id"])].add(module_id)

    rows: list[ReportRow] = []
    for item in items:
        item_id = _text(item.get("id"))
        state_id = _state_id(item)
        state = states_by_id.get(state_id, {})
        state_type = _text(state.get("type") or state.get("group"))
        if state_ids and state_id not in state_ids:
            continue
        if not normalized_filters["include_completed"] and state_type.casefold() in {"completed", "cancelled"}:
            continue
        item_modules = module_by_item_id.get(item_id, set())
        if module_ids and not (item_modules & module_ids):
            continue
        item_labels = _nested_ids(item.get("labels"))
        if label_ids and not label_ids.issubset(item_labels):
            continue
        assignee_ids = _nested_ids(item.get("assignees"), nested="member")
        if normalized_filters["assignee_ids"] and not set(normalized_filters["assignee_ids"]).issubset(assignee_ids):
            continue
        priority = _text(item.get("priority") or "none")
        if normalized_filters["priorities"] and priority.casefold() not in {value.casefold() for value in normalized_filters["priorities"]}:
            continue
        updated = item.get("updated_at")
        updated_at = _text(updated)
        parsed_updated = None
        if updated_at:
            try:
                parsed_updated = datetime.fromisoformat(updated_at.replace("Z", "+00:00"))
                if parsed_updated.tzinfo is None:
                    parsed_updated = parsed_updated.replace(tzinfo=timezone.utc)
            except ValueError:
                parsed_updated = None
        if normalized_filters["updated_after"] and (not parsed_updated or parsed_updated < normalized_filters["updated_after"]):
            continue
        if normalized_filters["updated_before"] and (not parsed_updated or parsed_updated > normalized_filters["updated_before"]):
            continue
        module_names = ", ".join(_text(modules_by_id[module_id].get("name")) for module_id in sorted(item_modules) if module_id in modules_by_id)
        label_names = ", ".join(_text(labels_by_id[label_id].get("name")) for label_id in sorted(item_labels) if label_id in labels_by_id)
        rows.append(
            ReportRow(
                identifier=_identifier(project, item),
                title=_text(item.get("name")),
                state=_text(state.get("name") or item.get("state")),
                state_type=state_type,
                priority=priority,
                assignees=_display_names(item.get("assignees"), nested="member"),
                updated_at=updated_at[:10],
                module=module_names,
                labels=label_names,
                description=_description(item),
            )
        )
    report_title = _text(title) or f"{_text(project.get('name') or project.get('identifier') or project_id)} work items"
    if len(report_title) > 180:
        raise ReportError("title must be 180 characters or fewer.")
    return ReportDocument(
        title=report_title,
        project=project,
        generated_at=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        filters={key: value.isoformat() if isinstance(value, datetime) else value for key, value in normalized_filters.items()},
        columns=list(normalized_layout["columns"]),
        group_by=normalized_layout["group_by"],
        rows=sorted(rows, key=lambda row: (row.state.casefold(), row.identifier.casefold())),
        total_work_items=total_count,
    )


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_table_widths(table: Any, widths: Iterable[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)


def _column_widths(columns: list[str]) -> list[float]:
    weights = {"identifier": 0.85, "title": 2.1, "state": 0.85, "priority": 0.6, "assignees": 0.95, "updated_at": 0.85, "module": 1.0, "labels": 1.1, "description": 2.3}
    selected = [weights[column] for column in columns]
    scale = 6.5 / sum(selected)
    return [value * scale for value in selected]


def _header_text(column: str) -> str:
    return {"identifier": "ID", "title": "Title", "state": "State", "priority": "Priority", "assignees": "Assignees", "updated_at": "Updated", "module": "Module", "labels": "Labels", "description": "Description"}[column]


def render_docx(report: ReportDocument, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(report.title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(14)
    metadata.add_run(f"{_text(report.project.get('name'))} | Generated {report.generated_at}").italic = True
    document.add_heading("Summary", level=1)
    total_text = f"Selected {len(report.rows)} work item(s)"
    if report.total_work_items is not None:
        total_text += f" from {report.total_work_items} in the project"
    document.add_paragraph(total_text + ".")
    summary = Counter(row.state or "Unspecified" for row in report.rows)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_table_widths(table, [4.6, 1.9])
    for cell, value in zip(table.rows[0].cells, ("State", "Items"), strict=True):
        cell.text = value
        _set_cell_shading(cell, "E8EEF5")
        cell.paragraphs[0].runs[0].bold = True
    for state, count in sorted(summary.items()):
        cells = table.add_row().cells
        cells[0].text = state
        cells[1].text = str(count)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("Work items", level=1)
    widths = _column_widths(report.columns)
    for group_name, rows in report.groups:
        if report.group_by != "none":
            document.add_heading(group_name, level=2)
        item_table = document.add_table(rows=1, cols=len(report.columns))
        item_table.style = "Table Grid"
        _set_table_widths(item_table, widths)
        header_cells = item_table.rows[0].cells
        for cell, column in zip(header_cells, report.columns, strict=True):
            cell.text = _header_text(column)
            _set_cell_shading(cell, "F2F4F7")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in rows:
            cells = item_table.add_row().cells
            for cell, column in zip(cells, report.columns, strict=True):
                cell.text = row.value(column) or "-"
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if column in {"identifier", "priority", "updated_at"} else WD_ALIGN_PARAGRAPH.LEFT
        document.add_paragraph()
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Plane Workflow report")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def render_pdf(report: ReportDocument, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=21, leading=25, textColor=colors.HexColor("#0B2545"), spaceAfter=6)
    metadata_style = ParagraphStyle("Metadata", parent=styles["Normal"], fontName="Helvetica-Oblique", fontSize=9, textColor=colors.HexColor("#4B5563"), spaceAfter=14)
    heading_style = ParagraphStyle("Heading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#2E74B5"), spaceBefore=14, spaceAfter=6)
    cell_style = ParagraphStyle("Cell", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.5, leading=9)
    header_style = ParagraphStyle("Header", parent=cell_style, fontName="Helvetica-Bold", alignment=TA_CENTER, textColor=colors.HexColor("#0B2545"))
    document = SimpleDocTemplate(str(output_path), pagesize=letter, leftMargin=0.55 * inch, rightMargin=0.55 * inch, topMargin=0.65 * inch, bottomMargin=0.65 * inch, title=report.title)
    story: list[Any] = [Paragraph(report.title, title_style), Paragraph(f"{_text(report.project.get('name'))} | Generated {report.generated_at}", metadata_style), Paragraph("Summary", heading_style)]
    total_text = f"Selected {len(report.rows)} work item(s)"
    if report.total_work_items is not None:
        total_text += f" from {report.total_work_items} in the project"
    story.extend([Paragraph(total_text + ".", styles["BodyText"]), Spacer(1, 8)])
    summary = [[Paragraph("State", header_style), Paragraph("Items", header_style)]] + [[Paragraph(state, cell_style), Paragraph(str(count), cell_style)] for state, count in sorted(Counter(row.state or "Unspecified" for row in report.rows).items())]
    summary_table = Table(summary, colWidths=[4.7 * inch, 1.2 * inch], repeatRows=1)
    summary_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")), ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8D0DA")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("ALIGN", (1, 1), (1, -1), "CENTER"), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    story.extend([summary_table, Paragraph("Work items", heading_style)])
    widths = [width / 6.5 * 7.4 * inch for width in _column_widths(report.columns)]
    for group_name, rows in report.groups:
        if report.group_by != "none":
            story.append(Paragraph(group_name, heading_style))
        data = [[Paragraph(_header_text(column), header_style) for column in report.columns]]
        data.extend([[Paragraph(row.value(column) or "-", cell_style) for column in report.columns] for row in rows])
        table = Table(data, colWidths=widths, repeatRows=1, splitByRow=True)
        table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CBD5E1")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
        story.extend([table, Spacer(1, 8)])
    document.build(story)


def export_work_items_report(
    api: Any,
    *,
    project_id: str,
    report_format: ReportFormat,
    title: str | None,
    filters: dict[str, Any] | None,
    layout: dict[str, Any] | None,
    output_directory: str | None,
) -> dict[str, Any]:
    if report_format not in {"docx", "pdf"}:
        raise ReportError("format must be docx or pdf.")
    report = build_report(api, project_id=project_id, title=title, filters=filters, layout=layout)
    directory = Path(output_directory).expanduser() if output_directory else Path.cwd() / "plane-reports"
    safe_identifier = re.sub(r"[^A-Za-z0-9._-]+", "-", _text(report.project.get("identifier") or project_id)).strip("-") or "project"
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    path = directory.resolve() / f"plane-work-items-{safe_identifier}-{timestamp}.{report_format}"
    if report_format == "docx":
        render_docx(report, path)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        render_pdf(report, path)
        mime_type = "application/pdf"
    return {
        "status": "exported",
        "path": str(path),
        "format": report_format,
        "mime_type": mime_type,
        "project": {"id": report.project.get("id"), "identifier": report.project.get("identifier"), "name": report.project.get("name")},
        "work_item_count": len(report.rows),
        "project_work_item_count": report.total_work_items,
        "filters": report.filters,
        "layout": {"group_by": report.group_by, "columns": report.columns},
        "note": "The report is read-only for Plane. It does not change any Plane work items.",
    }
